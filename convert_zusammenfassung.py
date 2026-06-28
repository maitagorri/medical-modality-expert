"""Convert docs/zusammenfassung.md → a tight one-page Word document."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_lines = Path("docs/zusammenfassung.md").read_text().splitlines()

doc = Document()

# --- Page layout: narrow margins ---
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# --- Base style: small font, tight line spacing ---
style = doc.styles["Normal"]
style.font.name = "DejaVu Sans"
style.font.size = Pt(10)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(4)
pf.line_spacing = 1.0


def add_runs(paragraph, text: str):
    """Parse inline **bold**, *italic*, and `code` into runs."""
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
        else:
            paragraph.add_run(tok)


for line in md_lines:
    line = line.rstrip()
    if not line:
        continue
    if line.startswith("# "):
        continue  # drop the title to save space
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_runs(p, line[2:])
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, line)

doc.save("zusammenfassung.docx")
print("Saved zusammenfassung.docx")
